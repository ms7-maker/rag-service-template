"""
Очистка локальных документов для RAG: .txt, .docx, .pdf.

Каждый файл → data/<имя>_clean.txt

Запуск из assistant_api:
  python cleaner.py about.docx
  python cleaner.py data/raw/catalog.pdf
  python cleaner.py                    # все файлы в data/raw/
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

from dotenv import load_dotenv

from data_paths import DATA_DIR, clean_output_path, stem_from_input_path

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
DEFAULT_INPUT = DATA_DIR / "raw"
MIN_TEXT_LEN = 100
SUPPORTED_SUFFIXES = {".txt", ".docx", ".pdf"}

_SPACE_CHARS = str.maketrans(
    {
        "\u00a0": " ",
        "\u2007": " ",
        "\u202f": " ",
        "\u200b": "",
        "\ufeff": "",
    }
)


def _resolve_path(raw: str) -> Path:
    p = Path(raw.strip().lstrip("/\\"))
    return p if p.is_absolute() else HERE / p


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "Для .docx установите: pip install python-docx"
        ) from exc

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("Для .pdf установите: pip install pypdf") from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def read_document(path: Path) -> str:
    """Извлечь сырой текст из файла."""
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(
            f"Формат {suffix} не поддерживается. Допустимо: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )
    if suffix == ".txt":
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    return ""


def _collect_inputs(path: Path) -> list[Path]:
    if not path.exists():
        raise FileNotFoundError(f"Не найдено: {path}")

    if path.is_file():
        return [path]

    files: list[Path] = []
    for ext in SUPPORTED_SUFFIXES:
        files.extend(path.glob(f"*{ext}"))
    files = sorted(set(files))
    if not files:
        raise FileNotFoundError(
            f"В папке нет .txt/.docx/.pdf: {path}"
        )
    return files


def _parse_input_paths(raw: str) -> list[Path]:
    parts = [p.strip() for p in raw.split(",") if p.strip()]
    if not parts:
        return _collect_inputs(DEFAULT_INPUT)

    out: list[Path] = []
    for part in parts:
        resolved = _resolve_path(part)
        if resolved.is_dir():
            out.extend(_collect_inputs(resolved))
        else:
            out.append(resolved)
    return out


def _resolve_cli_args(argv: list[str]) -> list[Path] | None:
    """Аргументы командной строки: python cleaner.py about.docx [file2 ...]"""
    if not argv:
        return None
    paths: list[Path] = []
    for arg in argv:
        p = _resolve_path(arg)
        if p.is_dir():
            paths.extend(_collect_inputs(p))
        else:
            paths.append(p)
    return paths


def clean_text(text: str) -> str:
    """Нормализация текста: пробелы, пустые строки, подряд идущие дубликаты."""
    text = text.translate(_SPACE_CHARS)
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines: list[str] = []
    prev: str | None = None

    for raw_line in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw_line.strip())
        if not line:
            if lines and lines[-1] != "":
                lines.append("")
            prev = None
            continue
        if line == prev:
            continue
        lines.append(line)
        prev = line

    result = "\n".join(lines)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Очистка документов (.txt, .docx, .pdf) -> data/<имя>_clean.txt",
    )
    parser.add_argument(
        "files",
        nargs="*",
        help="Файл или папка (например about.docx или data/raw/)",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    load_dotenv(ROOT / ".env")

    parser = _build_arg_parser()
    args = parser.parse_args(argv)
    cli_paths = _resolve_cli_args(args.files) if args.files else None

    try:
        if cli_paths is not None:
            inputs = cli_paths
        else:
            input_raw = os.getenv("CLEANER_INPUT_PATH", "").strip()
            if input_raw:
                inputs = _parse_input_paths(input_raw)
            else:
                inputs = _collect_inputs(DEFAULT_INPUT)
    except FileNotFoundError as exc:
        print(exc, file=sys.stderr)
        print(
            "Укажите файл: python cleaner.py about.docx",
            file=sys.stderr,
        )
        print(
            "или CLEANER_INPUT_PATH в .env, или положите файлы в data/raw/",
            file=sys.stderr,
        )
        return 1

    print(f"Файлов: {len(inputs)}")
    DATA_DIR.mkdir(parents=True, exist_ok=True)

    saved = 0
    for path in inputs:
        if not path.is_file():
            print(f"Пропуск (не файл): {path}", file=sys.stderr)
            continue
        try:
            raw = read_document(path)
        except (ValueError, ImportError, OSError) as exc:
            print(f"Ошибка {path.name}: {exc}", file=sys.stderr)
            continue

        cleaned = clean_text(raw)
        stem = stem_from_input_path(path)
        out_path = clean_output_path(stem)
        print(f"  {path.name}: {len(raw)} -> {len(cleaned)} символов")
        if not cleaned:
            print(f"     пропуск (пусто): {out_path.name}", file=sys.stderr)
            continue
        if len(cleaned) < MIN_TEXT_LEN:
            print(f"     предупреждение: мало текста в {out_path.name}", file=sys.stderr)
        out_path.write_text(cleaned, encoding="utf-8")
        print(f"     -> {out_path}")
        saved += 1

    if saved == 0:
        print("Нет текста для сохранения", file=sys.stderr)
        return 1

    print(f"\nГотово: сохранено файлов {saved} в {DATA_DIR}")
    print("Для пересборки индекса удалите папку faiss_db и снова запустите app.py")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
